"""Report generation service.

Generates DOCX, PDF, and XLSX reports from project answers.
Uses python-docx for DOCX, reportlab for PDF, and XlsxWriter for XLSX.
"""

from __future__ import annotations

import logging
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any

import asyncpg

from app.core.config import settings
from app.schemas.reports import BrandingConfig, ReportFormat, ReportType

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Generate DD reports in various formats."""

    def __init__(self, db: asyncpg.Connection):
        self.db = db

    async def generate(
        self,
        project_id: str,
        report_type: ReportType,
        output_format: ReportFormat,
        branding: BrandingConfig | None = None,
        category_filter: str | None = None,
        language: str = "de",
    ) -> Path:
        """Generate a report and return the output file path."""

        branding = branding or BrandingConfig()

        # Fetch project data
        project = await self._get_project(project_id)
        answers = await self._get_answers(project_id, report_type, category_filter)

        # Determine output path
        exports_dir = settings.projects_dir / project_id / "exports"
        exports_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{report_type.value}_{timestamp}.{output_format.value}"
        output_path = exports_dir / filename

        if output_format == ReportFormat.DOCX:
            await self._generate_docx(
                output_path, project, answers, report_type, branding, language
            )
        elif output_format == ReportFormat.PDF:
            await self._generate_pdf(
                output_path, project, answers, report_type, branding, language
            )
        elif output_format == ReportFormat.XLSX:
            await self._generate_xlsx(
                output_path, project, answers, report_type, language
            )

        logger.info("Report generated: %s", output_path)
        return output_path

    async def _get_project(self, project_id: str) -> dict[str, Any]:
        row = await self.db.fetchrow(
            "SELECT * FROM projects WHERE id = $1", project_id
        )
        return dict(row) if row else {}

    async def _get_answers(
        self,
        project_id: str,
        report_type: ReportType,
        category_filter: str | None,
    ) -> list[dict[str, Any]]:
        """Fetch answers with question metadata, filtered by report type."""
        query = """
            SELECT a.*, q.category, q.subcategory, q.question_de,
                   q.question_en, q.priority, q.severity_weight,
                   q.expected_format
            FROM answers a
            JOIN questions q ON a.question_id = q.id
            WHERE a.project_id = $1
        """
        params: list[Any] = [project_id]
        param_idx = 2

        if report_type == ReportType.RED_FLAGS:
            query += " AND (a.confidence_tier IN ('low', 'insufficient_data'))"
        if category_filter:
            query += f" AND q.category = ${param_idx}"
            params.append(category_filter)
            param_idx += 1

        query += " ORDER BY q.category, q.priority, q.id"

        rows = await self.db.fetch(query, *params)
        return [dict(r) for r in rows]

    async def _generate_docx(
        self,
        output_path: Path,
        project: dict,
        answers: list[dict],
        report_type: ReportType,
        branding: BrandingConfig,
        language: str,
    ) -> None:
        """Generate a DOCX report using python-docx."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()

        # Cover page
        title = doc.add_heading(level=0)
        title.text = project.get("name", "DD Report")

        doc.add_paragraph(
            f"Asset Class: {project.get('asset_class', 'N/A')}"
        )
        doc.add_paragraph(
            f"Date: {datetime.now().strftime('%d.%m.%Y')}"
        )
        doc.add_paragraph(
            f"Classification: {branding.classification}"
        )
        if branding.company_name:
            doc.add_paragraph(f"Prepared by: {branding.company_name}")

        doc.add_page_break()

        # Group answers by category
        categories: dict[str, list[dict]] = {}
        for a in answers:
            cat = a.get("category", "Other")
            categories.setdefault(cat, []).append(a)

        # Content
        for category, cat_answers in categories.items():
            doc.add_heading(category, level=1)

            for answer in cat_answers:
                q_key = "question_de" if language == "de" else "question_en"
                question_text = answer.get(q_key, answer.get("question_de", ""))

                # Priority badge
                priority = answer.get("priority", "medium")
                p = doc.add_heading(level=2)
                p.text = f"[{priority.upper()}] {question_text}"

                # Answer
                doc.add_paragraph(answer.get("answer_text", "No answer"))

                # Confidence
                tier = answer.get("confidence_tier", "unknown")
                score = answer.get("confidence_score", 0)
                doc.add_paragraph(
                    f"Confidence: {tier.upper()} ({score:.0%})",
                )

                doc.add_paragraph("")  # Spacing

        doc.save(str(output_path))

    async def _generate_pdf(
        self,
        output_path: Path,
        project: dict,
        answers: list[dict],
        report_type: ReportType,
        branding: BrandingConfig,
        language: str,
    ) -> None:
        """Generate a PDF report using reportlab."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        doc = SimpleDocTemplate(str(output_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story: list = []

        # Title
        story.append(
            Paragraph(project.get("name", "DD Report"), styles["Title"])
        )
        story.append(Spacer(1, 24))
        story.append(
            Paragraph(
                f"Asset Class: {project.get('asset_class', 'N/A')}<br/>"
                f"Date: {datetime.now().strftime('%d.%m.%Y')}<br/>"
                f"Classification: {branding.classification}",
                styles["Normal"],
            )
        )
        story.append(Spacer(1, 36))

        # Answers
        for answer in answers:
            q_key = "question_de" if language == "de" else "question_en"
            question_text = answer.get(q_key, "")

            story.append(
                Paragraph(f"<b>{question_text}</b>", styles["Heading3"])
            )
            story.append(
                Paragraph(answer.get("answer_text", "No answer"), styles["Normal"])
            )

            tier = answer.get("confidence_tier", "unknown")
            story.append(
                Paragraph(f"<i>Confidence: {tier.upper()}</i>", styles["Normal"])
            )
            story.append(Spacer(1, 12))

        doc.build(story)

    async def _generate_xlsx(
        self,
        output_path: Path,
        project: dict,
        answers: list[dict],
        report_type: ReportType,
        language: str,
    ) -> None:
        """Generate an XLSX Q&A matrix using XlsxWriter."""
        import xlsxwriter

        workbook = xlsxwriter.Workbook(str(output_path))
        worksheet = workbook.add_worksheet("Q&A Matrix")

        # Header format
        header_fmt = workbook.add_format({
            "bold": True,
            "bg_color": "#1a365d",
            "font_color": "#ffffff",
            "border": 1,
        })

        # Confidence color formats
        conf_formats = {
            "high": workbook.add_format({"bg_color": "#c6efce"}),
            "medium": workbook.add_format({"bg_color": "#ffeb9c"}),
            "low": workbook.add_format({"bg_color": "#ffc7ce"}),
            "insufficient_data": workbook.add_format({"bg_color": "#d9d9d9"}),
        }

        headers = [
            "Question ID", "Category", "Priority", "Question",
            "Answer", "Confidence", "Score", "Sources", "Status",
        ]
        for col, header in enumerate(headers):
            worksheet.write(0, col, header, header_fmt)

        q_key = "question_de" if language == "de" else "question_en"

        for row, answer in enumerate(answers, start=1):
            worksheet.write(row, 0, answer.get("question_id", ""))
            worksheet.write(row, 1, answer.get("category", ""))
            worksheet.write(row, 2, answer.get("priority", ""))
            worksheet.write(row, 3, answer.get(q_key, ""))
            worksheet.write(row, 4, answer.get("answer_text", ""))

            tier = answer.get("confidence_tier", "unknown")
            fmt = conf_formats.get(tier)
            worksheet.write(row, 5, tier.upper(), fmt)
            worksheet.write(row, 6, answer.get("confidence_score", 0))
            worksheet.write(row, 7, "")  # Sources populated separately
            worksheet.write(row, 8, answer.get("status", ""))

        worksheet.autofilter(0, 0, len(answers), len(headers) - 1)
        worksheet.set_column(3, 3, 50)  # Question column width
        worksheet.set_column(4, 4, 60)  # Answer column width

        workbook.close()
