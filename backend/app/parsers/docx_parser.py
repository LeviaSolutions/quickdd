"""DOCX parser using python-docx.

Extracts body text, headers, footers, comments, and tables.
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from app.parsers.base import (
    BaseParser,
    ExtractedPage,
    ExtractedTable,
    ParseResult,
)

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Parse Microsoft Word .docx files."""

    @property
    def supported_mimes(self) -> set[str]:
        return {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    async def parse(self, file_path: Path) -> ParseResult:
        from docx import Document

        result = ParseResult()

        try:
            doc = Document(str(file_path))
        except Exception as exc:
            logger.error("Failed to open DOCX %s: %s", file_path, exc)
            result.warnings.append(f"Failed to open: {exc}")
            return result

        # Extract core properties as metadata
        try:
            props = doc.core_properties
            result.metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else "",
                "last_modified_by": props.last_modified_by or "",
            }
        except Exception:
            result.metadata = {}

        # Extract paragraphs grouped into logical pages
        # (DOCX does not have strict page boundaries, so we treat
        # the entire document as a single logical page)
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract header/footer text
        header_texts: list[str] = []
        footer_texts: list[str] = []
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        header_texts.append(para.text.strip())
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        footer_texts.append(para.text.strip())

        full_text_parts: list[str] = []
        if header_texts:
            full_text_parts.append("[Headers]\n" + "\n".join(set(header_texts)))
        full_text_parts.append("\n".join(paragraphs))
        if footer_texts:
            full_text_parts.append("[Footers]\n" + "\n".join(set(footer_texts)))

        result.pages.append(
            ExtractedPage(
                page_number=1,
                text="\n\n".join(full_text_parts),
            )
        )
        result.page_count = 1

        # Extract tables
        for idx, table in enumerate(doc.tables):
            try:
                headers = [
                    cell.text.strip() for cell in table.rows[0].cells
                ] if table.rows else []

                rows = [
                    [cell.text.strip() for cell in row.cells]
                    for row in table.rows[1:]
                ]

                result.tables.append(
                    ExtractedTable(
                        table_index=idx,
                        headers=headers,
                        rows=rows,
                        page_number=1,
                    )
                )
            except Exception as exc:
                logger.warning("Table %d extraction failed: %s", idx, exc)
                result.warnings.append(f"Table {idx} extraction failed: {exc}")

        return result
